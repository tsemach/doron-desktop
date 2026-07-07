import sys
import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def txt_to_docx(txt_path, docx_path):
    doc = docx.Document()
    
    with open(txt_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    paragraphs = content.split('\n\n')
    for p in paragraphs:
        p_text = p.strip()
        if not p_text:
            continue
            
        is_heading = False
        heading_level = 1
        
        if p_text.startswith('###'):
            is_heading = True
            heading_level = 3
            p_text = p_text[3:].strip()
        elif p_text.startswith('##'):
            is_heading = True
            heading_level = 2
            p_text = p_text[2:].strip()
        elif p_text.startswith('#'):
            is_heading = True
            heading_level = 1
            p_text = p_text[1:].strip()
        elif p_text.startswith('===') or p_text.startswith('---'):
            continue
        elif len(p_text) < 80 and ('חוזה' in p_text or 'הסכם' in p_text or 'תצהיר' in p_text or 'צוואה' in p_text or 'מכתב' in p_text or 'דוח' in p_text or 'פסק דין' in p_text or 'סיכום' in p_text or 'רישיון' in p_text or 'בקשה' in p_text) and '\n' not in p_text:
            is_heading = True
            heading_level = 2
            
        if is_heading:
            h = doc.add_heading(level=heading_level)
            h.paragraph_format.right_to_left = True
            h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = h.add_run(p_text)
            run.font.rtl = True
        else:
            lines = p_text.split('\n')
            for line in lines:
                line_text = line.strip()
                if not line_text:
                    continue
                para = doc.add_paragraph()
                para.paragraph_format.right_to_left = True
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = para.add_run(line_text)
                run.font.rtl = True
                
    doc.save(docx_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python create_docx.py <txt_path> <docx_path>")
        sys.exit(1)
    txt_to_docx(sys.argv[1], sys.argv[2])
